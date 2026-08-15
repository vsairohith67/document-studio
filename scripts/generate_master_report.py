from __future__ import annotations

import csv
import re
from pathlib import Path
from PIL import Image, ImageOps, ImageEnhance
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK

ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / 'report'
REPORT_DIR.mkdir(parents=True, exist_ok=True)
ASSET_DIR = REPORT_DIR / '_assets'
ASSET_DIR.mkdir(exist_ok=True)
DOCX = REPORT_DIR / 'Document_Studio_Master_Blueprint.docx'

# ---------- image preparation ----------
image_sources = {
    'home': ROOT/'apps/prototype/screenshots/home.png',
    'workbench': ROOT/'apps/prototype/screenshots/workbench.png',
    'system': ROOT/'diagrams/system-architecture.png',
    'ia': ROOT/'diagrams/information-architecture.png',
    'job': ROOT/'diagrams/job-lifecycle.png',
    'router': ROOT/'diagrams/execution-router.png',
    'roadmap': ROOT/'diagrams/roadmap.png',
}
images = {}
for key, src in image_sources.items():
    im = Image.open(src).convert('L')
    im = ImageOps.autocontrast(im)
    im = ImageEnhance.Contrast(im).enhance(1.08)
    out = ASSET_DIR/f'{key}-gray.png'
    im.save(out, optimize=True)
    images[key] = out

# ---------- helpers ----------
def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'), 'dxa')


def set_table_borders(table, color='555555', size='5'):
    tblPr = table._tbl.tblPr
    borders = tblPr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tblPr.append(borders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        tag = f'w:{edge}'
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag); borders.append(element)
        element.set(qn('w:val'),'single'); element.set(qn('w:sz'),size); element.set(qn('w:color'),color)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'),'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'),'preserve'); instrText.text=' PAGE '
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'),'end')
    run._r.extend([fldChar1,instrText,fldChar2])


def add_code_block(doc, text):
    p = doc.add_paragraph(style='Code Block')
    p.paragraph_format.keep_together = True
    r = p.add_run(text.rstrip())
    r.font.name = 'Consolas'; r.font.size = Pt(7.5)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'),'F1F1F1'); pPr.append(shd)
    return p


def add_table(doc, rows, widths=None, font_size=8.0):
    if not rows:
        return None
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for j in range(cols):
        cell = hdr.cells[j]
        text = rows[0][j] if j < len(rows[0]) else ''
        cell.text = text
        set_cell_shading(cell, 'D9D9D9')
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.bold = True; run.font.size = Pt(font_size)
    for ri, row in enumerate(rows[1:], start=1):
        cells = table.add_row().cells
        for j in range(cols):
            cell = cells[j]
            cell.text = row[j] if j < len(row) else ''
            if ri % 2 == 0:
                set_cell_shading(cell, 'F6F6F6')
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs: run.font.size = Pt(font_size)
    if widths:
        for row in table.rows:
            for j, width in enumerate(widths[:cols]):
                row.cells[j].width = width
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def clean_inline(text):
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'\1 (\2)', text)
    return text.replace('  ', ' ')


def add_markdown(doc, md):
    lines = md.splitlines()
    i = 0
    para = []
    in_code = False
    code_lines = []

    def flush_para():
        nonlocal para
        if para:
            text = clean_inline(' '.join(x.strip() for x in para))
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.line_spacing = 1.08
            para = []

    while i < len(lines):
        line = lines[i]
        if line.strip().startswith('```'):
            flush_para()
            if not in_code:
                in_code = True; code_lines=[]
            else:
                add_code_block(doc, '\n'.join(code_lines)); in_code=False
            i += 1; continue
        if in_code:
            code_lines.append(line); i += 1; continue
        if not line.strip():
            flush_para(); i += 1; continue
        if line.startswith('# '):
            flush_para(); i += 1; continue  # file title already added
        if line.startswith('## '):
            flush_para(); doc.add_heading(clean_inline(line[3:]), level=2); i += 1; continue
        if line.startswith('### '):
            flush_para(); doc.add_heading(clean_inline(line[4:]), level=3); i += 1; continue
        if line.startswith('#### '):
            flush_para(); doc.add_heading(clean_inline(line[5:]), level=4); i += 1; continue
        if line.startswith('|') and i+1 < len(lines) and lines[i+1].startswith('|') and set(lines[i+1].replace('|','').replace('-','').replace(':','').strip()) == set():
            flush_para()
            rows=[]
            while i < len(lines) and lines[i].startswith('|'):
                parts=[clean_inline(x.strip()) for x in lines[i].strip().strip('|').split('|')]
                if parts and all(re.fullmatch(r':?-+:?', p) for p in parts):
                    i += 1; continue
                rows.append(parts); i += 1
            add_table(doc, rows, font_size=7.7)
            continue
        m = re.match(r'^[-*] \[([ xX])\] (.*)$', line)
        if m:
            flush_para(); marker = '[DONE] ' if m.group(1).lower() == 'x' else '[TODO] '
            p=doc.add_paragraph(style='List Bullet'); p.add_run(marker + clean_inline(m.group(2))); p.paragraph_format.space_after=Pt(2); i+=1; continue
        if re.match(r'^[-*] ', line):
            flush_para(); p=doc.add_paragraph(clean_inline(re.sub(r'^[-*] ','',line)), style='List Bullet'); p.paragraph_format.space_after=Pt(2); i+=1; continue
        if re.match(r'^\d+\. ', line):
            flush_para(); p=doc.add_paragraph(clean_inline(line)); p.paragraph_format.left_indent=Cm(.55); p.paragraph_format.first_line_indent=Cm(-.42); p.paragraph_format.space_after=Pt(2); i+=1; continue
        if line.startswith('> '):
            flush_para(); p=doc.add_paragraph(clean_inline(line[2:]), style='Quote'); i+=1; continue
        para.append(line); i += 1
    flush_para()
    if in_code and code_lines: add_code_block(doc,'\n'.join(code_lines))


def add_figure(doc, path, caption, width=Inches(6.8)):
    # Keep the artwork and its caption in one paragraph so the caption can never be orphaned.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run()
    shape = r.add_picture(str(path), width=width)
    # Screen-reader description; using the human-readable figure caption keeps the
    # generated DOCX accessible without changing its visual layout.
    shape._inline.docPr.set('descr', caption)
    shape._inline.docPr.set('title', caption)
    br = p.add_run(); br.add_break()
    cap = p.add_run(caption)
    cap.italic = True; cap.font.name = 'Aptos'; cap.font.size = Pt(8); cap.font.color.rgb = RGBColor(60,60,60)

# ---------- document ----------
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
sec.top_margin = Cm(1.7); sec.bottom_margin = Cm(1.6); sec.left_margin = Cm(1.7); sec.right_margin = Cm(1.7)
sec.header_distance = Cm(.7); sec.footer_distance = Cm(.7)
sec.different_first_page_header_footer = True

styles = doc.styles
styles['Normal'].font.name='Aptos'; styles['Normal'].font.size=Pt(9.2)
styles['Normal'].paragraph_format.space_after=Pt(5); styles['Normal'].paragraph_format.line_spacing=1.08
for name,size,bold in [('Title',30,True),('Heading 1',18,True),('Heading 2',13,True),('Heading 3',10.5,True),('Heading 4',9.5,True)]:
    st=styles[name]; st.font.name='Aptos Display' if name in ('Title','Heading 1','Heading 2') else 'Aptos'; st.font.size=Pt(size); st.font.bold=bold; st.font.color.rgb=RGBColor(0,0,0)
styles['Heading 1'].paragraph_format.space_before=Pt(14); styles['Heading 1'].paragraph_format.space_after=Pt(7); styles['Heading 1'].paragraph_format.keep_with_next=True
styles['Heading 2'].paragraph_format.space_before=Pt(10); styles['Heading 2'].paragraph_format.space_after=Pt(5); styles['Heading 2'].paragraph_format.keep_with_next=True
styles['Heading 3'].paragraph_format.space_before=Pt(7); styles['Heading 3'].paragraph_format.space_after=Pt(3); styles['Heading 3'].paragraph_format.keep_with_next=True
styles['Caption'].font.name='Aptos'; styles['Caption'].font.size=Pt(8); styles['Caption'].font.italic=True; styles['Caption'].font.color.rgb=RGBColor(60,60,60)
if 'Code Block' not in styles:
    st=styles.add_style('Code Block', WD_STYLE_TYPE.PARAGRAPH); st.font.name='Consolas'; st.font.size=Pt(7.5); st.paragraph_format.left_indent=Cm(.3); st.paragraph_format.right_indent=Cm(.3); st.paragraph_format.space_before=Pt(4); st.paragraph_format.space_after=Pt(6)

# Header/footer
header = sec.header.paragraphs[0]
header.text = 'DOCUMENT STUDIO - COMPLETE PRODUCT AND BUILD BLUEPRINT'
header.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in header.runs: r.font.size=Pt(8); r.font.bold=True; r.font.color.rgb=RGBColor(45,45,45)
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r=footer.add_run('Personal planning document  |  '); r.font.size=Pt(8); add_page_field(footer)

# Cover
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(75)
r=p.add_run('DOCUMENT STUDIO'); r.bold=True; r.font.name='Aptos Display'; r.font.size=Pt(34)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Complete Product, Architecture, UI/UX and Codex Build Blueprint'); r.font.size=Pt(17); r.bold=True
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(18)
r=p.add_run('Local-first document and PDF workspace'); r.font.size=Pt(13)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(28)
for line in ['Re-audited master specification','132 planned capabilities in one complete edition','Windows desktop first - optional web and mobile later','Black-and-white print edition']:
    q=doc.add_paragraph(); q.alignment=WD_ALIGN_PARAGRAPH.CENTER; q.add_run(line).font.size=Pt(10)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(60)
r=p.add_run('Prepared for V. Sai Rohith'); r.bold=True; r.font.size=Pt(11)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Version 2.0.1 - 22 July 2026').font.size=Pt(10)
doc.add_page_break()

# Audit summary
h=doc.add_heading('Re-audit outcome', level=1)
p=doc.add_paragraph('The earlier work was a strong starting point, but it was not ready to begin implementation without correction. The previous report still used the former name, several external-app updates were not verifiably complete, and the architecture/UI/model/dependency handoffs lacked enough operational detail. This master blueprint corrects those gaps and makes Document Studio the sole canonical name.')
rows=[
['Area','Final decision'],
['Product','Document Studio - one complete personal edition'],
['First platform','Windows desktop; macOS follows'],
['Desktop stack','Tauri 2.11.x + React 19.2.7+ + TypeScript + Vite 8.1.5+ + Rust'],
['Processing','Local by default; optional browser/cloud/external paths are explicit'],
['Viewer','PDF.js with progressive rendering and virtualized thumbnails'],
['Core engines','qpdf, libvips, OCRmyPDF/Tesseract, LibreOffice'],
['AI','Optional, provider-neutral, local-first, page-cited and consent-gated'],
['Design','Precision Paper - calm, professional, accessible and document-first'],
['Implementation','Codex vertical slices; Phase 0 before feature expansion'],
]
add_table(doc, rows, widths=[Inches(1.7), Inches(5.1)], font_size=8.5)
p=doc.add_paragraph('Planning package status: complete enough to create the private repository and start Phase 0. Live publication to Notion, Figma, GitHub or Canvs must be confirmed by their connectors; the import-ready assets are included in this package.')
p.paragraph_format.space_before=Pt(7)
doc.add_page_break()

# Contents
# Use a static contents list so the file is immediately printable without updating fields.
doc.add_heading('Contents', level=1)
contents=[
'1. Audit and completion record','2. Product charter','3. Personas and jobs to be done','4. Unified feature catalogue','5. Platform strategy','6. UI/UX strategy','7. Information architecture','8. Screen specifications','9. Performance architecture','10. System architecture','11. Unified operation contract','12. Data and database design','13. API and IPC design','14. Security, privacy and threat model','15. Dependency and license register','16. Hugging Face and local model plan','17. Test and quality strategy','18. Delivery, CI/CD and observability','19. Roadmap and milestones','20. Codex delivery method','21. Figma handoff','22. Notion knowledge-base plan','23. Implementation readiness checklist','24. Development setup and first build','25. Final recheck record','Appendix A. Full feature catalogue','Appendix B. Codex prompt pack','Appendix C. Primary implementation references'
]
for item in contents:
    p=doc.add_paragraph(item); p.paragraph_format.space_after=Pt(2)
doc.add_page_break()

# Assemble docs
ordered = [
'00-AUDIT-AND-COMPLETION.md','01-PRODUCT-CHARTER.md','02-PERSONAS-AND-JOBS.md','03-FEATURE-CATALOGUE.md','04-PLATFORM-STRATEGY.md','05-UX-STRATEGY.md','06-INFORMATION-ARCHITECTURE.md','07-SCREEN-SPECIFICATIONS.md','08-PERFORMANCE-ARCHITECTURE.md','09-SYSTEM-ARCHITECTURE.md','10-OPERATION-CONTRACT.md','11-DATA-AND-DATABASE-DESIGN.md','12-API-AND-IPC-DESIGN.md','13-SECURITY-PRIVACY-THREAT-MODEL.md','14-DEPENDENCY-AND-LICENSE-REGISTER.md','15-HUGGING-FACE-MODEL-PLAN.md','16-TEST-AND-QUALITY-STRATEGY.md','17-DELIVERY-CI-CD-OBSERVABILITY.md','18-ROADMAP-AND-MILESTONES.md','19-CODEX-DELIVERY-METHOD.md','20-FIGMA-HANDOFF.md','21-NOTION-KNOWLEDGE-BASE.md','22-IMPLEMENTATION-READINESS-CHECKLIST.md','23-DEVELOPMENT-SETUP.md','24-FINAL-RECHECK.md'
]
fig_after = {
'01-PRODUCT-CHARTER.md': ('roadmap','Figure 1. Product delivery sequence from foundation to optional ecosystem.', Inches(6.7)),
'05-UX-STRATEGY.md': ('home','Figure 2. Home-screen design direction: local status, drop zone, quick tools, recent work and saved workflows.', Inches(6.8)),
'06-INFORMATION-ARCHITECTURE.md': ('ia','Figure 3. Main information architecture.', Inches(6.5)),
'07-SCREEN-SPECIFICATIONS.md': ('workbench','Figure 4. Merge workbench concept with files/pages, document canvas, inspector and persistent job tray.', Inches(6.8)),
'08-PERFORMANCE-ARCHITECTURE.md': ('router','Figure 5. Execution-path router for desktop, browser-local and optional cloud work.', Inches(6.3)),
'09-SYSTEM-ARCHITECTURE.md': ('system','Figure 6. Canonical desktop system architecture.', Inches(6.7)),
'10-OPERATION-CONTRACT.md': ('job','Figure 7. Verified job lifecycle and alternate terminal states.', Inches(6.7)),
'18-ROADMAP-AND-MILESTONES.md': ('roadmap','Figure 8. Milestone roadmap.', Inches(6.7)),
}
for idx, fn in enumerate(ordered, start=1):
    md=(ROOT/'docs'/fn).read_text(encoding='utf-8')
    title=md.splitlines()[0].lstrip('# ').strip()
    heading = doc.add_heading(f'{idx}. {title}', level=1)
    if idx > 1:
        heading.paragraph_format.page_break_before = True
    add_markdown(doc, md)
    if fn in fig_after:
        key, cap, width=fig_after[fn]
        add_figure(doc, images[key], cap, width)

# Appendix A full feature catalogue
# Landscape section for readability
sec2=doc.add_section(WD_SECTION_START.NEW_PAGE)
sec2.orientation = 1  # landscape
sec2.page_width = Cm(29.7); sec2.page_height = Cm(21.0)
sec2.top_margin=Cm(1.4); sec2.bottom_margin=Cm(1.4); sec2.left_margin=Cm(1.3); sec2.right_margin=Cm(1.3)
# link header/footer to previous
sec2.header.is_linked_to_previous=True; sec2.footer.is_linked_to_previous=True

doc.add_heading('Appendix A. Full unified feature catalogue', level=1)
p=doc.add_paragraph('The following 132 capabilities are planned as one complete product. Phase expresses sequence, not a subscription tier. “Must/Should/Could” is delivery priority.')
with (ROOT/'docs/feature-catalog.csv').open(encoding='utf-8') as f:
    feats=list(csv.DictReader(f))
rows=[['Category','Feature','Required behavior','Phase / priority / engine']]
for x in feats:
    rows.append([x['category'],x['feature'],x['description'],f"P{x['phase']} · {x['priority']} · {x['engine']}"])
add_table(doc, rows, widths=[Cm(3.0),Cm(4.5),Cm(11.8),Cm(7.0)], font_size=6.8)

# Back to portrait
sec3=doc.add_section(WD_SECTION_START.NEW_PAGE)
sec3.orientation=0
sec3.page_width=Cm(21.0); sec3.page_height=Cm(29.7)
sec3.top_margin=Cm(1.7);sec3.bottom_margin=Cm(1.6);sec3.left_margin=Cm(1.7);sec3.right_margin=Cm(1.7)
sec3.header.is_linked_to_previous=True; sec3.footer.is_linked_to_previous=True

doc.add_heading('Appendix B. Codex prompt pack', level=1)
for fn in sorted((ROOT/'codex/prompts').glob('*.md')):
    md=fn.read_text(encoding='utf-8')
    title=md.splitlines()[0].lstrip('# ').strip()
    doc.add_heading(title, level=2)
    add_markdown(doc, md)

# references
refs=(ROOT/'docs/REFERENCES.md').read_text(encoding='utf-8')
doc.add_page_break(); doc.add_heading('Appendix C. Primary implementation references', level=1); add_markdown(doc, refs)

# Package inventory
p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(12)
r=p.add_run('Companion package contents'); r.bold=True
for line in ['Codex-ready repository scaffold','Interactive static UI prototype and two reference screens','Figma tokens, file structure, components and plugin plan','Hugging Face/local-model registry','Architecture and workflow diagrams in DOT, Mermaid, PNG and SVG','Notion import pack with source research and historical reports','Validation scripts, CI skeleton, ADRs and acceptance checklists']:
    doc.add_paragraph(line, style='List Bullet')

# Metadata
props=doc.core_properties
props.title='Document Studio - Complete Product, Architecture, UI/UX and Codex Build Blueprint'
props.subject='Re-audited master software blueprint'
props.author='V. Sai Rohith / Document Studio planning'
props.keywords='Document Studio, PDF, Tauri, React, Rust, Codex, UI UX, architecture'
props.comments='Black-and-white print edition; generated 22 July 2026.'

doc.save(DOCX)
print(DOCX)
